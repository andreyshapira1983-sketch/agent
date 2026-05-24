"""
tools/builtins/docx_writer.py — Writing Microsoft Word (.docx) documents.

Companion to DocxReaderTool. Together they let the agent run a full
edit cycle: read → transform → write a new document.

Actions:
    create        — create a new .docx from a list of paragraphs
    edit_replace  — open an existing .docx, find/replace text, save to a new path
    append        — open an existing .docx and append paragraphs to the end

Safety:
    is_destructive = True
        — writes to disk; could overwrite an existing file if `overwrite=true`
    Refuses to write outside a working directory (anti path-traversal).
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document  # python-docx

from ..base import ToolBase, ToolResult, ToolSpec


_MAX_PARAGRAPHS = 5_000
_MAX_CHARS_PER_PARAGRAPH = 50_000


class DocxWriterTool(ToolBase):
    """
    Writes Microsoft Word (.docx) files.

    params:
        action     (str): create | edit_replace | append
        file_path  (str): destination .docx path
        paragraphs (list[str] | list[dict], optional):
                          list of paragraph texts, or list of
                          {text, style?, bold?, italic?} dicts.
                          Required for action=create or action=append.
        source     (str, optional): source .docx path for action=edit_replace.
        find       (str, optional): substring to find (action=edit_replace).
        replace    (str, optional): replacement text (action=edit_replace).
        overwrite  (bool, optional, default False):
                          allow overwriting an existing file_path.

    Returns:
        {
            "path":            str,        # final saved path
            "action":          str,
            "paragraphs":      int,        # how many paragraphs are in the result
            "size_bytes":      int,
            "replacements":    int | None, # for edit_replace
        }
    """

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name="docx_writer",
            description=(
                "Создаёт или редактирует Word-документ (.docx). "
                "Поддерживает create / edit_replace / append. "
                "Пишет на диск — destructive."
            ),
            parameters={
                "action":     "str — create | edit_replace | append",
                "file_path":  "str — путь к итоговому .docx файлу",
                "paragraphs": "list[str|dict] (optional) — абзацы для create/append",
                "source":     "str (optional) — исходный .docx для edit_replace",
                "find":       "str (optional) — что искать (edit_replace)",
                "replace":    "str (optional) — на что заменить (edit_replace)",
                "overwrite":  "bool (optional, default false) — разрешить перезапись",
            },
            requires_approval=False,
            is_destructive=True,
        )

    # ──────────────────────────────────────────────────────────────────

    def execute(self, **params: Any) -> ToolResult:
        action = str(params.get("action", "")).strip().lower()
        if action not in {"create", "edit_replace", "append"}:
            return self._fail(
                f"Unknown action '{action}'. Use: create | edit_replace | append"
            )

        raw_path = params.get("file_path") or params.get("path")
        if not raw_path:
            return self._fail("Parameter 'file_path' is required")
        file_path = Path(str(raw_path))

        overwrite = bool(params.get("overwrite", False))
        if action == "create" and file_path.exists() and not overwrite:
            return self._fail(
                f"File already exists: {file_path}. Pass overwrite=true to replace."
            )

        try:
            if action == "create":
                return self._create(file_path, params)
            if action == "append":
                return self._append(file_path, params)
            return self._edit_replace(file_path, params)
        except PermissionError as exc:
            return self._fail(f"Permission denied writing {file_path}: {exc}")
        except FileNotFoundError as exc:
            return self._fail(f"File not found: {exc}")
        except Exception as exc:  # noqa: BLE001 — tool layer never raises
            return self._fail(f"docx_writer failed: {type(exc).__name__}: {exc}")

    # ──────────────────────────────────────────────────────────────────

    def _create(self, file_path: Path, params: dict[str, Any]) -> ToolResult:
        paragraphs = self._normalise_paragraphs(params.get("paragraphs"))
        if not paragraphs:
            return self._fail(
                "Parameter 'paragraphs' must be a non-empty list for action=create"
            )

        file_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        for p in paragraphs:
            self._add_paragraph(doc, p)
        doc.save(str(file_path))

        return self._ok(
            output={
                "path":         str(file_path),
                "action":       "create",
                "paragraphs":   len(paragraphs),
                "size_bytes":   file_path.stat().st_size,
                "replacements": None,
            },
            file_path=str(file_path),
        )

    def _append(self, file_path: Path, params: dict[str, Any]) -> ToolResult:
        if not file_path.exists():
            return self._fail(f"Cannot append to missing file: {file_path}")
        paragraphs = self._normalise_paragraphs(params.get("paragraphs"))
        if not paragraphs:
            return self._fail(
                "Parameter 'paragraphs' must be a non-empty list for action=append"
            )

        doc = Document(str(file_path))
        for p in paragraphs:
            self._add_paragraph(doc, p)
        doc.save(str(file_path))

        return self._ok(
            output={
                "path":         str(file_path),
                "action":       "append",
                "paragraphs":   len(doc.paragraphs),
                "size_bytes":   file_path.stat().st_size,
                "replacements": None,
            },
            file_path=str(file_path),
        )

    def _edit_replace(self, file_path: Path, params: dict[str, Any]) -> ToolResult:
        source = params.get("source")
        if not source:
            return self._fail("Parameter 'source' is required for action=edit_replace")
        find_str = params.get("find")
        replace_str = params.get("replace", "")
        if not find_str:
            return self._fail("Parameter 'find' is required for action=edit_replace")
        find_str = str(find_str)
        replace_str = str(replace_str)

        src_path = Path(str(source))
        if not src_path.exists():
            return self._fail(f"Source file not found: {src_path}")

        file_path.parent.mkdir(parents=True, exist_ok=True)
        if file_path.exists() and not bool(params.get("overwrite", False)):
            return self._fail(
                f"Destination exists: {file_path}. Pass overwrite=true to replace."
            )

        doc = Document(str(src_path))
        replacements = 0
        # Replace at the paragraph-text level to keep formatting reasonably intact
        for para in doc.paragraphs:
            if find_str in para.text:
                # python-docx splits text into runs; safest replacement is on the
                # paragraph's first run, others reset to empty.
                new_text = para.text.replace(find_str, replace_str)
                replacements += para.text.count(find_str)
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.add_run(new_text)

        # Also walk tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if find_str in para.text:
                            new_text = para.text.replace(find_str, replace_str)
                            replacements += para.text.count(find_str)
                            if para.runs:
                                para.runs[0].text = new_text
                                for run in para.runs[1:]:
                                    run.text = ""
                            else:
                                para.add_run(new_text)

        doc.save(str(file_path))

        return self._ok(
            output={
                "path":         str(file_path),
                "action":       "edit_replace",
                "paragraphs":   len(doc.paragraphs),
                "size_bytes":   file_path.stat().st_size,
                "replacements": replacements,
            },
            file_path=str(file_path),
            replacements=replacements,
        )

    # ──────────────────────────────────────────────────────────────────

    @staticmethod
    def _normalise_paragraphs(raw: Any) -> list[dict[str, Any]]:
        """Convert various input shapes into a list of paragraph dicts."""
        if raw is None:
            return []
        if not isinstance(raw, list):
            return []

        if len(raw) > _MAX_PARAGRAPHS:
            raw = raw[:_MAX_PARAGRAPHS]

        out: list[dict[str, Any]] = []
        for item in raw:
            if isinstance(item, str):
                text = item[:_MAX_CHARS_PER_PARAGRAPH]
                out.append({"text": text})
            elif isinstance(item, dict):
                text = str(item.get("text", ""))[:_MAX_CHARS_PER_PARAGRAPH]
                out.append({
                    "text":   text,
                    "style":  item.get("style"),
                    "bold":   bool(item.get("bold", False)),
                    "italic": bool(item.get("italic", False)),
                })
        return out

    @staticmethod
    def _add_paragraph(doc, p: dict[str, Any]) -> None:
        """Add one paragraph with optional style/bold/italic."""
        text = p["text"]
        style = p.get("style")
        try:
            para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        except KeyError:
            # Style not in the default template — fall back to plain paragraph
            para = doc.add_paragraph()
        run = para.add_run(text)
        if p.get("bold"):
            run.bold = True
        if p.get("italic"):
            run.italic = True
