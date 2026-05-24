"""
tools/builtins/office_reader.py — Чтение документов Microsoft Office

DocxReaderTool  — .docx (Word)
XlsxReaderTool  — .xlsx (Excel)

Использует python-docx и openpyxl (чистый Python, без LibreOffice/COM).

DocxReaderTool actions:
    extract_text   — весь текст документа (параграфы + таблицы)
    get_metadata   — автор, название, дата создания, страницы, символы
    list_headings  — список заголовков (структура документа)

XlsxReaderTool actions:
    extract_data   — данные листа в виде CSV-подобных строк
    get_sheets     — список листов книги
    get_metadata   — автор, название, дата, листы, размер
"""

from __future__ import annotations

import time
from pathlib import Path
from typing import Any

from ..base import ToolBase, ToolResult, ToolSpec

_MAX_FILE_SIZE_MB = 50
_MAX_CHARS_OUTPUT = 100_000


# ══════════════════════════════════════════════════════════════════════
#  DocxReaderTool
# ══════════════════════════════════════════════════════════════════════

class DocxReaderTool(ToolBase):
    """
    Читает Word-документы (.docx).

    params:
        file_path  (str): Путь к .docx файлу
        action     (str, optional): extract_text | get_metadata | list_headings
                                    (по умолчанию: extract_text)
        max_chars  (int, optional): Максимум символов (по умолчанию 20000)
    """

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name="docx_reader",
            description="Читает текст и метаданные из Word-документов (.docx)",
            parameters={
                "file_path": "str — путь к .docx файлу",
                "action":    "str (optional) — extract_text | get_metadata | list_headings",
                "max_chars": "int (optional, default 20000) — максимум символов",
            },
            requires_approval=False,
            is_destructive=False,
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            import docx  # noqa: F401
        except ImportError:
            return self._fail("python-docx не установлен. Запустите: pip install python-docx")

        file_path = params.get("file_path", "")
        if not file_path:
            return self._fail("Параметр 'file_path' обязателен")

        path = Path(str(file_path))
        if not path.exists():
            return self._fail(f"Файл не найден: {file_path}")
        if not path.is_file():
            return self._fail(f"Это не файл: {file_path}")
        if path.suffix.lower() != ".docx":
            return self._fail(f"Ожидается .docx, получено: {path.suffix!r}")

        size_mb = path.stat().st_size / (1024 * 1024)
        if size_mb > _MAX_FILE_SIZE_MB:
            return self._fail(f"Файл слишком большой: {size_mb:.1f} МБ (макс. {_MAX_FILE_SIZE_MB})")

        action = params.get("action", "extract_text")
        t0 = time.perf_counter()

        try:
            import docx

            doc = docx.Document(str(path))

            if action == "get_metadata":
                props = doc.core_properties
                result = {
                    "title":       props.title or "",
                    "author":      props.author or "",
                    "subject":     props.subject or "",
                    "created":     str(props.created) if props.created else "",
                    "modified":    str(props.modified) if props.modified else "",
                    "paragraphs":  len(doc.paragraphs),
                    "tables":      len(doc.tables),
                    "file_size_kb": round(path.stat().st_size / 1024, 1),
                }
                elapsed = (time.perf_counter() - t0) * 1000
                return self._ok(output=result, duration_ms=round(elapsed, 2))

            elif action == "list_headings":
                headings: list[dict] = []
                for para in doc.paragraphs:
                    if para.style.name.startswith("Heading"):
                        level = para.style.name.split()[-1]
                        try:
                            level_int = int(level)
                        except ValueError:
                            level_int = 1
                        headings.append({"level": level_int, "text": para.text.strip()})
                elapsed = (time.perf_counter() - t0) * 1000
                return self._ok(output=headings, duration_ms=round(elapsed, 2))

            elif action == "extract_text":
                max_chars = int(params.get("max_chars", 20_000))
                max_chars = min(max(100, max_chars), _MAX_CHARS_OUTPUT)

                parts: list[str] = []

                # Параграфы
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)

                # Таблицы
                for i, table in enumerate(doc.tables):
                    parts.append(f"\n[Таблица {i + 1}]")
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        parts.append(" | ".join(cells))

                text = "\n".join(parts)
                truncated = False
                if len(text) > max_chars:
                    text = text[:max_chars]
                    truncated = True

                elapsed = (time.perf_counter() - t0) * 1000
                return self._ok(
                    output=text,
                    duration_ms=round(elapsed, 2),
                    paragraphs=len(doc.paragraphs),
                    tables=len(doc.tables),
                    chars_returned=len(text),
                    truncated=truncated,
                )

            else:
                return self._fail(
                    f"Неизвестное действие: {action!r}. "
                    "Допустимо: extract_text | get_metadata | list_headings"
                )

        except Exception as exc:
            return self._fail(f"Ошибка чтения DOCX: {exc}")


# ══════════════════════════════════════════════════════════════════════
#  XlsxReaderTool
# ══════════════════════════════════════════════════════════════════════

class XlsxReaderTool(ToolBase):
    """
    Читает Excel-таблицы (.xlsx).

    params:
        file_path   (str): Путь к .xlsx файлу
        action      (str, optional): extract_data | get_sheets | get_metadata
                                     (по умолчанию: extract_data)
        sheet       (str | int, optional): Имя или индекс листа (с 0). По умолчанию — первый.
        max_rows    (int, optional): Максимум строк (по умолчанию 500)
        max_cols    (int, optional): Максимум колонок (по умолчанию 50)
    """

    @property
    def spec(self) -> ToolSpec:
        return ToolSpec(
            name="xlsx_reader",
            description="Читает данные и метаданные из Excel-файлов (.xlsx)",
            parameters={
                "file_path": "str — путь к .xlsx файлу",
                "action":    "str (optional) — extract_data | get_sheets | get_metadata",
                "sheet":     "str или int (optional) — имя или индекс листа (с 0)",
                "max_rows":  "int (optional, default 500) — максимум строк для извлечения",
                "max_cols":  "int (optional, default 50) — максимум колонок",
            },
            requires_approval=False,
            is_destructive=False,
        )

    def execute(self, **params: Any) -> ToolResult:
        try:
            import openpyxl  # noqa: F401
        except ImportError:
            return self._fail("openpyxl не установлен. Запустите: pip install openpyxl")

        file_path = params.get("file_path", "")
        if not file_path:
            return self._fail("Параметр 'file_path' обязателен")

        path = Path(str(file_path))
        if not path.exists():
            return self._fail(f"Файл не найден: {file_path}")
        if not path.is_file():
            return self._fail(f"Это не файл: {file_path}")
        if path.suffix.lower() not in {".xlsx", ".xlsm"}:
            return self._fail(f"Ожидается .xlsx / .xlsm, получено: {path.suffix!r}")

        size_mb = path.stat().st_size / (1024 * 1024)
        if size_mb > _MAX_FILE_SIZE_MB:
            return self._fail(f"Файл слишком большой: {size_mb:.1f} МБ (макс. {_MAX_FILE_SIZE_MB})")

        action = params.get("action", "extract_data")
        t0 = time.perf_counter()

        try:
            import openpyxl

            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)

            if action == "get_sheets":
                sheets = wb.sheetnames
                elapsed = (time.perf_counter() - t0) * 1000
                return self._ok(output=sheets, duration_ms=round(elapsed, 2))

            elif action == "get_metadata":
                props = wb.properties
                result = {
                    "title":        props.title or "",
                    "creator":      props.creator or "",
                    "description":  props.description or "",
                    "created":      str(props.created) if props.created else "",
                    "modified":     str(props.modified) if props.modified else "",
                    "sheets":       wb.sheetnames,
                    "sheet_count":  len(wb.sheetnames),
                    "file_size_kb": round(path.stat().st_size / 1024, 1),
                }
                elapsed = (time.perf_counter() - t0) * 1000
                return self._ok(output=result, duration_ms=round(elapsed, 2))

            elif action == "extract_data":
                max_rows = int(params.get("max_rows", 500))
                max_cols = int(params.get("max_cols", 50))
                max_rows = min(max(1, max_rows), 5000)
                max_cols = min(max(1, max_cols), 500)

                # Выбор листа
                sheet_param = params.get("sheet")
                if sheet_param is None:
                    ws = wb.active
                elif isinstance(sheet_param, int):
                    name = wb.sheetnames[sheet_param]
                    ws = wb[name]
                else:
                    if str(sheet_param) not in wb.sheetnames:
                        return self._fail(
                            f"Лист {sheet_param!r} не найден. "
                            f"Доступные: {wb.sheetnames}"
                        )
                    ws = wb[str(sheet_param)]

                rows_data: list[list[str]] = []
                for row_idx, row in enumerate(ws.iter_rows(max_row=max_rows, max_col=max_cols)):
                    row_values = [
                        str(cell.value) if cell.value is not None else ""
                        for cell in row
                    ]
                    # Пропускаем полностью пустые строки
                    if any(v.strip() for v in row_values):
                        rows_data.append(row_values)

                # Форматируем как текст (CSV-стиль)
                lines = [" | ".join(row) for row in rows_data]
                text = "\n".join(lines)

                elapsed = (time.perf_counter() - t0) * 1000
                sheet_name = ws.title if hasattr(ws, "title") else str(sheet_param)
                return self._ok(
                    output=text,
                    duration_ms=round(elapsed, 2),
                    sheet=sheet_name,
                    rows_returned=len(rows_data),
                    truncated=ws.max_row is not None and ws.max_row > max_rows,
                )

            else:
                return self._fail(
                    f"Неизвестное действие: {action!r}. "
                    "Допустимо: extract_data | get_sheets | get_metadata"
                )

        except Exception as exc:
            return self._fail(f"Ошибка чтения XLSX: {exc}")
